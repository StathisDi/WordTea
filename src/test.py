#! python3

import sys
import argparse
import os
import re
import comtypes.client
import csv
import docx
import WordTea_functions
import WordTea_classes
from WordTea_functions import *
from docx import Document
from docx.shared import Inches
from WordTea_classes import reference_list


# Main function


def __main__():
    print("main function")
    ###########################################################
    # Use parser to get the cmd arguments #####################
    ###########################################################
    parser = argparse.ArgumentParser(
        description='WordTea: Word document parser. Generates citations and cross-references from latex like text in word document files.')
    parser.add_argument('inFile', help='Path to the input file.')
    parser.add_argument('pdfFile', help='Path to the output pdf file.')
    parser.add_argument('tmpFile', help='Path to the output temporary file.')
    parser.add_argument(
        '-r', metavar='ref', help='Path to the reference file (if a reference file is to be used).')
    parser.add_argument(
        '-i', metavar='iter', help='Max iterations used when searching to complete a cross-reference tag, default 5.', default=5)
    parser.add_argument('--silent', action='store_false',
                        help='Disable Verbose level 1, basic status print for missed references and citations inside the document')
    parser.add_argument('--verbose', action='store_true',
                        help='Enable Verbose level 2, extreme error print for script debug')
    parser.add_argument('-s1', metavar='s1_format',
                        help='Format of the section 1 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-s2', metavar='s2_format',
                        help='Format of the section 2 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-s3', metavar='s3_format',
                        help='Format of the section 3 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-table', metavar='table_format',
                        help='Format of the table reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)

    args = parser.parse_args()

    ############################################################

    # Constant used for the pdf creation
    wdFormatPDF = 17

    #################################################
    # Old way of reading args from command ##########
    #################################################
    #in_file = os.path.abspath(sys.argv[1])
    #out_file = os.path.abspath(sys.argv[2])
    #tmp_file = os.path.abspath(sys.argv[3])
    #################################################
    # Reading the arguments of the command to #######
    # the variables used in the script        #######
    #################################################
    in_file = os.path.abspath(args.inFile)
    out_file = os.path.abspath(args.pdfFile)
    tmp_file = os.path.abspath(args.tmpFile)
    v1 = args.silent
    v2 = args.verbose
    max = args.i
    s1_f = int(args.s1)
    s2_f = int(args.s2)
    s3_f = int(args.s3)
    tbl_f = int(args.table)
    # Error hnadling
    if not isinstance(s1_f, type(1)):
        raise TypeError("-s1 expected integer, got %s" % type(input))
    if not isinstance(s2_f, type(1)):
        raise TypeError("-s2 expected integer, got %s" % type(input))
    if not isinstance(s3_f, type(1)):
        raise TypeError("-s3 expected integer, got %s" % type(input))
    if not isinstance(tbl_f, type(1)):
        raise TypeError("-s3 expected integer, got %s" % type(input))
    if not 0 < s1_f < 5:
        raise ValueError("-s1 argument must be between 1 and 4")
    if not 0 < s2_f < 5:
        raise ValueError("-s2 Argument must be between 1 and 4")
    if not 0 < s3_f < 5:
        raise ValueError("-s3 Argument must be between 1 and 4")
    if not 0 < tbl_f < 5:
        raise ValueError("-s1 argument must be between 1 and 4")

    if v2:
        print(max)

    ref_file = ""
    if (args.r) != None:
        ref_file = os.path.abspath(args.r)

    #######################################
    # Reading the main document
    document = Document(in_file)
    # Create the temporary document
    document.save(tmp_file)
    # Copy the main document to the temporary document
    document = Document(tmp_file)

    print(reference_list.__doc__)
    print(reference_list.__init__.__doc__)
    fig = reference_list("Figure", "f", "fig")
    sec0 = reference_list("Section 0", "s0", "sect0")
    sec1 = reference_list("Section 1", "s1", "sect1", 1, sec0)
    tbl = reference_list("Table", "t", "tbl")
    for pr in document.paragraphs:
        #print("Before :")
        # print(pr.text)
        fig.buildList(pr, True, True)
        sec0.buildList(pr, True, True)
        sec1.buildList(pr, True, True)
        tbl.buildList(pr, True, True)

        #print("After :")
        # print(pr.text)
    fig.printList()
    sec0.printList()
    sec1.printList()
    tbl.printList()
    fig.printParentList()
    sec0.printParentList()
    sec1.printParentList()
    tbl.printParentList()
    # exit()
    # for pr in document.paragraphs:
    #    print("Replace Before :")
    #    print(pr.text)
    #    parent_fig.matchNreplace(pr, True, True)
    #    parent_fig.matchNreplace(pr, True, True)
    #    parent_fig.matchNreplace(pr, True, True)
    #    print("Replace After :")
    #    print(pr.text)

    document.save(tmp_file)

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(tmp_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    return 1


if __name__ == "__main__":
    __main__()
