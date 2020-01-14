#! python3

# Copyright 2017, Dimitrios Stathis, All rights reserved.
# email         : stathis@kth.se, sta.dimitris@gmail.com
# Version       : 0.2.0
# Last edited   : 13/01/2020

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                         #
#This file is part of WordTea.                                            #
#                                                                         #
#    WordTea is free software: you can redistribute it and/or modify      #
#    it under the terms of the GNU General Public License as published by #
#    the Free Software Foundation, either version 3 of the License, or    #
#    (at your option) any later version.                                  #
#                                                                         #
#    WordTea is distributed in the hope that it will be useful,           #
#    but WITHOUT ANY WARRANTY; without even the implied warranty of       #
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the        #
#    GNU General Public License for more details.                         #
#                                                                         #
#    You should have received a copy of the GNU General Public License    #
#    along with WordTea.  If not, see <https://www.gnu.org/licenses/>.    #
#                                                                         #
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#

# This is a script for applying latex like references into a
# microsoft word document (.docx)


import sys
import argparse
import os
import re
import comtypes.client
import csv
import docx
import referenceList
from referenceList import referenceList
from docx import Document
from docx.shared import Inches

# TODO: There is bug that removes footnotes.
# TODO add functionality for reference from bibtex file


def __main__():
    ###########################################################
    # Use parser to get the cmd arguments #####################
    ###########################################################
    parser = argparse.ArgumentParser(description='WordTea: Word document parser. Generates citations and cross-references from latex like text in word document files.')
    parser.add_argument('inFile', help='Path to the input file.')
    parser.add_argument('pdfFile', help='Path to the output pdf file.')
    parser.add_argument('tmpFile', help='Path to the output temporary file.')
    parser.add_argument('-r', metavar='ref', help='Path to the reference file (if a reference file is to be used). !!!Not supported!!!')
    parser.add_argument('--silent', action='store_false', help='Disable Verbose level 1, basic status print for missed references and citations inside the document')
    parser.add_argument('--verbose', action='store_true', help='Enable Verbose level 2, extreme error print for script debug')
    parser.add_argument('-s1', metavar='s1_format', help='Format of the section 1 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-s2', metavar='s2_format', help='Format of the section 2 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-s3', metavar='s3_format', help='Format of the section 3 reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)
    parser.add_argument('-table', metavar='table_format', help='Format of the table reference style. Use 1 for normal numbering, 2 for Latin, 3 for small letter, 4 for capital letter, default 1.', default=1)

    args = parser.parse_args()

    ############################################################

    # Constant used for the pdf creation
    wdFormatPDF = 17

    #################################################
    # Reading the arguments of the command to #######
    # the variables used in the script        #######
    #################################################
    in_file = os.path.abspath(args.inFile)
    out_file = os.path.abspath(args.pdfFile)
    tmp_file = os.path.abspath(args.tmpFile)
    v1 = args.silent
    v2 = args.verbose
    s1_f = int(args.s1)
    s2_f = int(args.s2)
    s3_f = int(args.s3)
    tbl_f = int(args.table)
    # Error hnadling
    if not isinstance(s1_f, type(1)):
        raise TypeError("-s1 expected integer, got %s" % type(input))
    if not isinstance(s2_f, type(1)):
        raise TypeError("-s2 expected integer, got %s" % type(input))
    if not isinstance(s3_f, type(1)):
        raise TypeError("-s3 expected integer, got %s" % type(input))
    if not isinstance(tbl_f, type(1)):
        raise TypeError("-s3 expected integer, got %s" % type(input))
    if not 0 < s1_f < 5:
        raise ValueError("-s1 argument must be between 1 and 4")
    if not 0 < s2_f < 5:
        raise ValueError("-s2 Argument must be between 1 and 4")
    if not 0 < s3_f < 5:
        raise ValueError("-s3 Argument must be between 1 and 4")
    if not 0 < tbl_f < 5:
        raise ValueError("-s1 argument must be between 1 and 4")

    if v2:
        print(max)

    #ref_file = ""
    # if (args.r) != None:
    #    ref_file = os.path.abspath(args.r)

    #######################################
    print("Process: Open Document " + str(in_file) + ".")
    # Reading the main document
    document = Document(in_file)
    # Create the temporary document
    document.save(tmp_file)
    # Copy the main document to the temporary document
    document = Document(tmp_file)
    print("Process: File open.")

    ###############################################################
    # Create the reference classes ################################
    ###############################################################
    fig = referenceList("Figures", "fig", "fig", 1, None)
    sec1 = referenceList("Heading 1", "sec1", "sec1", s1_f)
    sec2 = referenceList("Heading 2", "sec2", "sec2", s2_f, sec1)
    sec3 = referenceList("Heading 3", "sec3", "sec3", s3_f, sec2)
    sec4 = referenceList("Heading 4", "sec4", "sec4", 1, sec3)
    tbl = referenceList("Tables", "tbl", "tbl", tbl_f, None)
    equ = referenceList("Equations", "eq", "eq", 1, None)
    cite = referenceList("Citations", "cite", "cite", 1, None)
    ###############################################################
    #
    print("Process: File open.")
    if v1:
        print("Details: Cross-reference database:")
        print(fig)
        print(sec1)
        print(sec2)
        print(sec3)
        print(sec4)
        print(tbl)
        print(equ)
        print(cite)

    ###############################################################
    # Build Reference database ####################################
    ###############################################################
    print("Process: Start building the reference database.")
    for pr in document.paragraphs:
        fig.buildList(pr, v1, v2)
        sec1.buildList(pr, v1, v2)
        sec2.buildList(pr, v1, v2)
        sec3.buildList(pr, v1, v2)
        sec4.buildList(pr, v1, v2)
        tbl.buildList(pr, v1, v2)
        equ.buildList(pr, v1, v2)
        cite.buildList(pr, v1, v2)
    print("Process: Building of the reference database is completed.")
    ###############################################################

    ###############################################################
    # Print List and other details ################################
    ###############################################################
    if v1:
        print("Details of the Database")
        fig.printList()
        sec1.printList()
        sec2.printList()
        tbl.printList()
        cite.printList()
    if v2:
        fig.printIndexList()
        sec1.printIndexList()
        sec2.printIndexList()
        tbl.printIndexList()
        cite.printIndexList()
    if v2:
        fig.printParentList()
        sec1.printParentList()
        sec2.printParentList()
        tbl.printParentList()
        cite.printParentList()
    ###############################################################

    ###############################################################
    # Search and Replace Labels ###################################
    ###############################################################
    print("Process: Start searching the text for relevant lables.")
    for pr in document.paragraphs:
        fig.matchNreplace(pr, v1, v2)
        sec1.matchNreplace(pr, v1, v2)
        sec2.matchNreplace(pr, v1, v2)
        sec3.matchNreplace(pr, v1, v2)
        sec4.matchNreplace(pr, v1, v2)
        tbl.matchNreplace(pr, v1, v2)
        equ.matchNreplace(pr, v1, v2)
        cite.matchNreplace(pr, v1, v2)
    print("Process: Search and replace of the labels in the text completed.")
    ###############################################################

    ###############################################################
    # Print the results from the search and replace function ######
    ###############################################################
    fig.checkRefList()
    sec1.checkRefList()
    sec2.checkRefList()
    sec3.checkRefList()
    sec3.checkRefList()
    tbl.checkRefList()
    equ.checkRefList()
    cite.checkRefList()
    ###############################################################
    if v1:
        print("Process: Cleaning up work space and deleting database.")
    ###############################################################
    # Delete and clear all objects ################################
    ###############################################################
    del fig
    del sec1
    del sec2
    del sec3
    del sec4
    del tbl
    del equ
    del cite
    ###############################################################
    print("Process: Saving post-processed temporary .docx document and .pdf file.")
    word = comtypes.client.CreateObject('Word.Application')
    # TODO Find how to use catch problem
    document.save(tmp_file)
    doc = word.Documents.Open(tmp_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    # os.remove(tmp_file)


if __name__ == "__main__":
    __main__()
    print("Execution successfully completed!")
