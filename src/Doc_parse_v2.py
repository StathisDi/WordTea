#! python3

# Copyright 2017, Dimitrios Stathis, All rights reserved.
# email         : stathis@kth
# Last edited   : 5/5/2018

#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#
#                                                                         #
#This file is part of DocParser.                                          #
#                                                                         #
#    DocParser is free software: you can redistribute it and/or modify    #
#    it under the terms of the GNU General Public License as published by #
#    the Free Software Foundation, either version 3 of the License, or    #
#    (at your option) any later version.                                  #
#                                                                         #
#    DocParser is distributed in the hope that it will be useful,         #
#    but WITHOUT ANY WARRANTY; without even the implied warranty of       #
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the        #
#    GNU General Public License for more details.                         #
#                                                                         #
#    You should have received a copy of the GNU General Public License    #
#    along with DocParser.  If not, see <https://www.gnu.org/licenses/>.  #
#                                                                         #
#~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~#

#This is a script for applying latex like references into a
#microsoft word document (.docx)


import sys
import argparse
import os
import re
import comtypes.client
import csv
import docx
from docx import Document
from docx.shared import Inches




parser = argparse.ArgumentParser(description='Word document parser. Generates citations and cross-references from latex like text in word document files.')
parser.add_argument('inFile', help='Path to the input file.')
parser.add_argument('pdfFile', help='Path to the output pdf file.')
parser.add_argument('tmpFile', help='Path to the output temporary file.')
parser.add_argument('-r', metavar='ref', help='Path to the reference file (if a reference file is to be used).')
parser.add_argument('-i', metavar='iter', help='Max iterations used when searching to complete a cross-reference tag, default 4.', default=4)
parser.add_argument('--silient', action='store_false', help='Disable Verbose level 1, basic status print for missed references and citations inside the document')
parser.add_argument('--verbose', action='store_true', help='Enable Verbose level 2, extreme error print for script debug')

args = parser.parse_args()



#TODO There is still a bug, tool might find broken down references and use the first one that it matches with
#           bug example:
#                       whatever_1 and whatever
#                       In this case the tool might use the reference whatever instead of whatever_1 since it will the first one that it finds
#
#           @ possible solutions for a fix:
#             Define a set of specific characters that are acceptable after a reference, for example accepted characters can be:
#              <space> . , [ ] etc but not characters like _ or any letter.
#TODO add functionality for reference from bibtex file


###########################################################
# Lists used for the different types of cross-references ##
###########################################################

# These lists are used when no reference file is used

section_0  = list()  # main sections

section_1  = list()
father_0   = list()  # father main section of the subsection

section_2  = list()
father_2_0 = list()  # father main section of the subsection
father_2_1 = list()  # father subsection (lvl 1)

section_3  = list()
father_3_0 = list()  # father main section of the subsection
father_3_1 = list()  # father subsection (lvl 1)
father_3_2 = list()  # father subsection (lvl 2)

figures     = list()
tables      = list()
citations   = list()
############################################################

# These lists are used when a reference file is used

ref_list_source = list()
ref_list_replace = list()

############################################################

# Constant used for the pdf creation
wdFormatPDF = 17

#################################################
# Old way of reading args from command ##########
#################################################
#in_file = os.path.abspath(sys.argv[1])
#out_file = os.path.abspath(sys.argv[2])
#tmp_file = os.path.abspath(sys.argv[3])
#################################################
# Reading the arguments of the command to #######
# the variables used in the script        #######
#################################################
in_file = os.path.abspath(args.inFile)
out_file = os.path.abspath(args.pdfFile)
tmp_file = os.path.abspath(args.tmpFile)
v1 = args.silient
v2 = args.verbose
max = args.i


if v2:
    print(max)

ref_file = ""
if (args.r) != None:
    ref_file = os.path.abspath(args.r)

#######################################
# Reading the main document
document = Document(in_file)
# Create the temporary document
document.save(tmp_file)
# Copy the main document to the temporary document
document = Document(tmp_file)


###############################################################
# Use the inline reference style without any external file ####
###############################################################
if ref_file == "":
    if v1:
        print('No reference file, searching for refs inside the document')
    #######################################
    # Build Database with references ######
    #######################################
    sec_0 = 0
    sec_1 = 0
    sec_2 = 0
    for pr in document.paragraphs:
        # lvl 0 sections
        if '^sec1'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'sec1'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    section_0.append(tmp1[1])
            pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
            sec_0 = sec_0 + 1
        # lvl 1 sections
        if '^sec2'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'sec2'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    section_1.append(tmp1[1])
                    father_0.append(sec_0)
                    pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
            sec_1 = sec_1 + 1
        # lvl 2 sections
        if '^sec3'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'sec3'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    section_2.append(tmp1[1])
                    father_2_0.append(sec_0)
                    father_2_1.append(sec_1)
                    pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
            sec_2 = sec_2 + 1
        # lvl 3 sections
        if '^sec4'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'sec4'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    section_1.append(tmp1[1])
                    father_3_0.append(sec_0)
                    father_3_1.append(sec_1)
                    father_3_2.append(sec_2)
            pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
        # figures
        if '^fig'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'fig{'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    if v2:
                        print(tmp1)
                    figures.append(tmp1[1])
            pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
        # tables
        if '^tbl'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'tbl'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    tables.append(tmp1[1])
            pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)
        # citations
        if '^cite'.lower() in pr.text.lower():
            # Loop added to work with runs (strings with same style)
            tmp = re.split('\^', pr.text)
            for pt in tmp:
                if 'cite{'.lower() in pt.lower():
                    remove_str = pt
                    tmp1 = re.split('{|}', pt)
                    if v2:
                        print(tmp1)
                    citations.append(tmp1[1])
            pattern = re.compile(re.escape(str(remove_str)), re.IGNORECASE)
            pr.text = pattern.sub('', pr.text)
            pr.text = re.sub('\^', '', pr.text)

    #####################################################
    # If Verbose level 1 is enable print the database ###
    #####################################################
    if v1:
        print('Sections lvl 1')
        print(section_0)
        print('Subsections lvl 2')
        print(section_1)
        print('Subsections lvl 3')
        print(section_2)
        print('Subsections lvl 4')
        print(section_3)
        print('Figures')
        print(figures)
        print('Tables')
        print(tables)
        print('Citations')
        print(citations)
        print('Number of Citations:')
        print(len(citations))

    #########################
    # search and replace ####
    #########################
    for p in document.paragraphs:
        inline = p.runs
        # Loop added to work with runs (strings with same style)
        for j in range(len(inline)):
            #print(inline[j].text)
            # sections lvl 0
            for i in range(len(section_0)):
                if str(section_0[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('ref' + str(section_0[i])), re.IGNORECASE)
                    inline[j].text = pattern.sub(str(i + 1), inline[j].text)

            # sections lvl 1
            for i in range(len(section_1)):
                if str(section_1[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('ref' + str(section_1[i])), re.IGNORECASE)
                    txt = str(father_0[i]+1) + '.' + str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

            # sections lvl 2
            for i in range(len(section_2)):
                if str(section_2[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('ref' + str(section_2[i])), re.IGNORECASE)
                    txt = str(father_2_0[i]+1) + '.' + str(father_2_1[i]+1) + '.' + str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

            # sections lvl 3
            for i in range(len(section_3)):
                if str(section_3[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('ref' + str(section_3[i])), re.IGNORECASE)
                    txt = str(father_3_0[i]+1) + '.' + str(father_3_1[i]+1) + '.' + str(father_3_2[i]+1) + '.' + str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

            # figures
            for i in range(len(figures)):
                if str(figures[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('fig' + str(figures[i])), re.IGNORECASE)
                    txt = str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

            # tables
            for i in range(len(tables)):
                if str(tables[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('ref' + str(tables[i])), re.IGNORECASE)
                    txt = str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

            # citations
            for i in range(len(citations)):
                if str(citations[i]).lower() in inline[j].text.lower():
                    #print(inline[j].text)
                    pattern = re.compile(re.escape('cit' + str(citations[i])), re.IGNORECASE)
                    txt = str(1 + i)
                    inline[j].text = pattern.sub(txt, inline[j].text)

    ################################################
    # Second run to fix misses, removes format #####
    ################################################
    # The second run tries to catch any kind of cross references that the first one missed.
    # It has one issue that depending on the order that the cross-refs are applied it might miss
    # and create a broken reference.
    # This bug needs to be fixed. See the beginning of the code for possible solutions @TODO
    if v1:
        print("The following citations/cross-refs were not caught by the normal search.")

    for p in document.paragraphs:
        for i in range(len(citations)):
            if str("cit"+str(citations[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("cite :")
                    print(citations[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "cit" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("cit" + str(citations[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                #print("IF "+str((j+k)==(len(inline))-1))
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    #print("length "+str(len(inline)))
                                    #print("j+k "+str(j+k))
                                    #print(flag)
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('cit' + str(citations[i])), re.IGNORECASE)
                        txt = str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)


        for i in range(len(figures)):
            if str("ref"+str(figures[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("figure :")
                    print(figures[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("ref" + str(figures[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(figures[i])), re.IGNORECASE)
                        txt = str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)

        for i in range(len(tables)):
            if str("ref"+str(tables[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("table :")
                    print(tables[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("ref" + str(tables[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(tables[i])), re.IGNORECASE)
                        txt = str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)

        # section lvl 0
        for i in range(len(section_0)):
            if str("ref"+str(section_0[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("Section lvl 1 :")
                    print(section_0[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("ref" + str(section_0[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(section_0[i])), re.IGNORECASE)
                        txt = str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)

        # section lvl 1
        for i in range(len(section_1)):
            if str("ref"+str(section_1[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("Section lvl 2 :")
                    print(section_1[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("ref" + str(section_1[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(section_1[i])), re.IGNORECASE)
                        txt = str(father_0[i]+1) + '.' + str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)

        # section lvl 2
        for i in range(len(section_2)):
            if str("ref"+str(section_2[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("Section lvl 3 :")
                    print(section_2[i])

                #print(p.text)
                for j in range(len(inline)):
                    #print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt=inline[j].text.lower()
                        k=1
                        flag=True
                        while flag:
                            #print(k)
                            if str("ref" + str(section_2[i].lower())) in tmp_txt:
                                flag=False
                            else:
                                if ((k>max) or ((j+k)==(len(inline))-1) or ((j+k)>=(len(inline)))):
                                    flag=False
                                else:
                                    tmp_txt=tmp_txt+inline[j+k].text.lower()
                                    inline[j].text=inline[j].text+inline[j+k].text
                                    inline[j+k].text = ""

                                k=k+1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(section_2[i])), re.IGNORECASE)
                        txt = str(father_2_0[i]+1) + '.' + str(father_2_1[i]+1) + '.' + str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)

        # section lvl 3
        for i in range(len(section_3)):
            if str("ref" + str(section_3[i].lower())) in p.text.lower():
                inline = p.runs
                if v1:
                    print("Section lvl 4 :")
                    print(section_3[i])

                # print(p.text)
                for j in range(len(inline)):
                    # print(inline[j].text)
                    if "ref" in inline[j].text.lower():
                        tmp_txt = inline[j].text.lower()
                        k = 1
                        flag = True
                        while flag:
                            # print(k)
                            if str("ref" + str(section_3[i].lower())) in tmp_txt:
                                flag = False
                            else:
                                if ((k > max) or ((j + k) == (len(inline)) - 1) or (
                                            (j + k) >= (len(inline)))):
                                    flag = False
                                else:
                                    tmp_txt = tmp_txt + inline[j + k].text.lower()
                                    inline[j].text = inline[j].text + inline[j + k].text
                                    inline[j + k].text = ""

                                k = k + 1
                        if v2:
                            print(tmp_txt)
                            print(inline[j].text)

                        pattern = re.compile(re.escape('ref' + str(section_3[i])), re.IGNORECASE)
                        txt = str(father_3_0[i]+1) + '.' + str(father_3_1[i]+1) + '.' + str(father_3_2[i]+1) + '.' + str(1 + i)
                        inline[j].text = pattern.sub(txt, inline[j].text)


########################################################################
# Use with reference file ##############################################
########################################################################
else:
    if v1:
        print('Using the reference file to parse the document')

    # read the ref .csv file
    with open(ref_file, newline='') as csvfile:
        spamreader = csv.reader(csvfile, delimiter=',')
        for row in spamreader:
            if len(row) != 0:
                ref_list_source.append(row[0])
                ref_list_replace.append(row[1])
    print(ref_list_replace)
    print(ref_list_source)
    # Go through the document and replace
    for p in document.paragraphs:
        for i in range(len(ref_list_source)):
            # if ref_list_source[i] in p.text:
            #    print(p.text)
            #    print(ref_list_source[i])
            #    p.text = re.sub(r + ref_list_source[i]),  ref_list_replace[i], r + p.text)
            # p.text = p.text.replace(str(ref_list_source[i]), ref_list_replace[i])
            # inline = p.runs
            if str(ref_list_source[i]).lower() in p.text.lower():
                print(p.text)
                print(str(ref_list_source[i]))
                pattern = re.compile(re.escape(str(ref_list_source[i])), re.IGNORECASE)
                p.text = pattern.sub(str(ref_list_replace[i]), p.text)

document.save(tmp_file)

word = comtypes.client.CreateObject('Word.Application')
doc = word.Documents.Open(tmp_file)
doc.SaveAs(out_file, FileFormat=wdFormatPDF)
doc.Close()
word.Quit()
# os.remove(tmp_file)
